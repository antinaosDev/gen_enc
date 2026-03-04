import docx
import sys

def extract_text(file_path):
    doc = docx.Document(file_path)
    full_text = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))
                
    return "\n".join(full_text)

if __name__ == "__main__":
    if len(sys.argv) > 1:
        print(extract_text(sys.argv[1]))
    else:
        print("Usage: python extract_docx.py <file_path>")
